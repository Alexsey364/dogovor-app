# -*- coding: utf-8 -*-
"""
Извлечение текста из загруженных файлов для полнотекстового поиска
и текстовых правил проверки.

Работает с PDF (текстовый слой) и DOCX. Старый .doc и сканы без
текстового слоя не читаются — для них текст остаётся пустым
(там понадобилось бы распознавание изображений, отдельный шаг).
"""

from __future__ import annotations

import os


def extract(path: str, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(path)
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
            return "\n".join(parts).strip()
        if ext == ".docx":
            import docx
            d = docx.Document(path)
            paras = [p.text for p in d.paragraphs]
            for table in d.tables:
                for row in table.rows:
                    paras.append(" ".join(cell.text for cell in row.cells))
            return "\n".join(paras).strip()
    except Exception:  # noqa: BLE001
        return ""
    return ""

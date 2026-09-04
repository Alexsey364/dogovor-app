# -*- coding: utf-8 -*-
"""
Конструктор договоров: собирает готовый .docx из реквизитов наших
компаний и заказчика по выбранному виду работ.

Тексты — переработанные на основе исторических форм, с доработкой в
нашу пользу (предоплата и её возврат, приёмка и молчаливое согласие,
гарантия и её отсчёт, неустойка, односторонний отказ, подсудность по
месту Исполнителя). Это ПРОЕКТ договора — перед применением его
утверждает юрист.
"""

from __future__ import annotations

import io
import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


MONTHS_RU = ["", "января", "февраля", "марта", "апреля", "мая", "июня",
             "июля", "августа", "сентября", "октября", "ноября", "декабря"]

KIND_TITLES = {
    "poverka": "поверки приборов учёта тепловой энергии",
    "montazh_uute": "монтажа узла учёта тепловой энергии",
    "obsluzhivanie": "технического обслуживания узла учёта",
    "postavka": "поставки оборудования",
    "proekt": "разработки проектной документации",
}


def _fmt_date(d):
    if not d:
        d = datetime.date.today()
    if isinstance(d, str):
        try:
            d = datetime.date.fromisoformat(d)
        except ValueError:
            return d
    return f"«{d.day:02d}» {MONTHS_RU[d.month]} {d.year} г."


def _money(v):
    try:
        return f"{float(v):,.2f}".replace(",", " ").replace(".", ",") + " руб."
    except (TypeError, ValueError):
        return "________ руб."


def _req_block(doc, party_label, org):
    """Реквизитный блок одной стороны."""
    doc.add_paragraph().add_run(party_label).bold = True
    lines = [org.get("name") or "—"]
    if org.get("address"):
        lines.append(f"Юридический адрес: {org['address']}")
    inn_kpp = "ИНН " + (org.get("inn") or "—")
    if org.get("kpp"):
        inn_kpp += f" КПП {org['kpp']}"
    lines.append(inn_kpp)
    if org.get("ogrn"):
        lines.append(f"ОГРН {org['ogrn']}")
    if org.get("bank_account"):
        lines.append(f"р/с {org['bank_account']}"
                     + (f" в {org['bank_name']}" if org.get("bank_name") else ""))
    if org.get("corr_account"):
        lines.append(f"к/с {org['corr_account']}"
                     + (f" БИК {org['bank_bik']}" if org.get("bank_bik") else ""))
    if org.get("phone"):
        lines.append(f"тел.: {org['phone']}")
    if org.get("email"):
        lines.append(f"e-mail: {org['email']}")
    for ln in lines:
        p = doc.add_paragraph(ln)
        p.paragraph_format.space_after = Pt(0)


def _h(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    p.paragraph_format.space_before = Pt(8)
    return p


def build(kind, number, sign_date, owner, cp, terms):
    """Собирает .docx договора. Возвращает BytesIO.

    owner, cp — словари реквизитов; terms — параметры (amount, advance_pct,
    warranty_months, deadline_days, subject).
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = KIND_TITLES.get(kind, "выполнения работ")
    owner_sign = owner.get("signatory") or "Индивидуальный предприниматель"
    owner_short = owner.get("short_name") or owner.get("name") or "Исполнитель"
    cp_sign = cp.get("signatory") or "руководитель"
    cp_name = cp.get("name") or "________"
    amount = terms.get("amount")
    advance = terms.get("advance_pct")
    warranty = terms.get("warranty_months") or 30
    deadline = terms.get("deadline_days") or 30
    subject = (terms.get("subject") or "").strip()

    # Заголовок
    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ph.add_run(f"ДОГОВОР № {number}\nна оказание услуг {title}")
    r.bold = True
    r.font.size = Pt(13)

    row = doc.add_paragraph()
    row.add_run("г. Уссурийск").bold = False
    row.add_run("\t\t\t\t\t\t" + _fmt_date(sign_date))

    # Преамбула
    pre = doc.add_paragraph()
    pre.add_run(f"{owner.get('name') or owner_short}, именуемый(ое) в дальнейшем "
                f"«Исполнитель», в лице {owner_sign} "
                f"{owner.get('director') or ''}, действующего на основании "
                f"{'свидетельства о регистрации' if 'предприниматель' in owner_sign.lower() else 'Устава'}, "
                f"с одной стороны, и {cp_name}, именуемое в дальнейшем «Заказчик», "
                f"в лице {cp_sign} {cp.get('director') or ''}, действующего на "
                f"основании {'Устава' if cp.get('kpp') else 'документа'}, с другой "
                f"стороны, заключили настоящий Договор о нижеследующем:")

    # 1. Предмет
    _h(doc, "1. Предмет Договора")
    doc.add_paragraph(
        f"1.1. Исполнитель обязуется оказать услуги {title}"
        + (f" ({subject})" if subject else "")
        + ", а Заказчик обязуется принять и оплатить оказанные услуги в порядке "
          "и на условиях, предусмотренных настоящим Договором.")
    doc.add_paragraph(
        "1.2. Объём, характеристики и место оказания услуг определяются "
        "Приложениями (спецификацией) к настоящему Договору, являющимися его "
        "неотъемлемой частью.")

    # 2. Цена и порядок расчётов (в нашу пользу — предоплата)
    _h(doc, "2. Цена Договора и порядок расчётов")
    doc.add_paragraph(
        f"2.1. Общая стоимость услуг по настоящему Договору составляет "
        f"{_money(amount)}, НДС не облагается на основании ст. 145 Налогового "
        f"кодекса РФ (либо НДС 5% — в зависимости от применяемого режима).")
    if advance and float(advance) >= 100:
        doc.add_paragraph(
            "2.2. Оплата производится Заказчиком в размере 100% (предоплата) в "
            "течение 5 (пяти) рабочих дней с момента выставления счёта и до "
            "начала оказания услуг. Исполнитель приступает к оказанию услуг "
            "после поступления оплаты на расчётный счёт Исполнителя.")
    else:
        pct = int(float(advance)) if advance else 50
        doc.add_paragraph(
            f"2.2. Заказчик выплачивает аванс в размере {pct}% в течение 5 (пяти) "
            f"рабочих дней с даты выставления счёта; окончательный расчёт — в "
            f"течение 5 (пяти) рабочих дней с даты подписания акта оказанных услуг.")
    doc.add_paragraph(
        "2.3. Обязанность Заказчика по оплате считается исполненной с момента "
        "поступления денежных средств на расчётный счёт Исполнителя. Все "
        "банковские расходы по перечислению несёт Заказчик.")

    # 3. Порядок сдачи-приёмки (молчаливое согласие в нашу пользу)
    _h(doc, "3. Порядок сдачи и приёмки услуг")
    doc.add_paragraph(
        "3.1. По завершении оказания услуг Исполнитель передаёт Заказчику акт "
        "оказанных услуг. Заказчик обязан в течение 5 (пяти) рабочих дней "
        "подписать акт либо направить мотивированный письменный отказ.")
    doc.add_paragraph(
        "3.2. Если в указанный срок Заказчик не подписал акт и не направил "
        "мотивированный отказ, услуги считаются принятыми в полном объёме и "
        "подлежат оплате, а акт — подписанным в одностороннем порядке.")

    # 4. Гарантия
    _h(doc, "4. Гарантийные обязательства")
    doc.add_paragraph(
        f"4.1. Гарантийный срок на результат оказанных услуг составляет "
        f"{warranty} календарных дней и исчисляется с даты подписания акта "
        f"оказанных услуг (для монтажных работ — с даты ввода узла учёта в "
        f"эксплуатацию).")
    doc.add_paragraph(
        "4.2. Гарантия не распространяется на недостатки, возникшие вследствие "
        "нарушения Заказчиком правил эксплуатации, вмешательства третьих лиц, "
        "аварий в сетях и иных обстоятельств, не зависящих от Исполнителя.")

    # 5. Срок
    _h(doc, "5. Сроки оказания услуг")
    doc.add_paragraph(
        f"5.1. Срок оказания услуг — до {deadline} рабочих дней, исчисляется с "
        f"даты поступления оплаты (аванса) на расчётный счёт Исполнителя.")

    # 6. Ответственность (в нашу пользу)
    _h(doc, "6. Ответственность Сторон")
    doc.add_paragraph(
        "6.1. За нарушение срока оплаты Заказчик уплачивает Исполнителю "
        "неустойку в размере 0,1% от неоплаченной суммы за каждый день "
        "просрочки, но не более 10% от суммы Договора.")
    doc.add_paragraph(
        "6.2. Совокупная ответственность Исполнителя по настоящему Договору "
        "ограничивается стоимостью фактически оказанных услуг. Исполнитель не "
        "возмещает упущенную выгоду и косвенные убытки.")

    # 7. Прочие условия
    _h(doc, "7. Прочие условия")
    doc.add_paragraph(
        "7.1. Исполнитель вправе в одностороннем внесудебном порядке отказаться "
        "от исполнения Договора при просрочке оплаты Заказчиком свыше 10 "
        "календарных дней, уведомив Заказчика письменно.")
    doc.add_paragraph(
        "7.2. Все споры разрешаются путём переговоров, при недостижении согласия "
        "— в Арбитражном суде Приморского края (по месту нахождения Исполнителя). "
        "Претензионный порядок обязателен, срок ответа на претензию — 10 "
        "календарных дней.")
    doc.add_paragraph(
        f"7.3. Договор вступает в силу с даты подписания и действует до "
        f"31.12.{(_yr(sign_date))} г., а в части взаиморасчётов — до полного "
        f"исполнения обязательств. Составлен в двух экземплярах, по одному для "
        f"каждой Стороны.")

    # 8. Реквизиты
    _h(doc, "8. Адреса и реквизиты Сторон")
    _req_block(doc, "ИСПОЛНИТЕЛЬ:", owner)
    doc.add_paragraph()
    _req_block(doc, "ЗАКАЗЧИК:", cp)
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run(f"Исполнитель: _______________ / {_short_fio(owner.get('director'))}\t\t"
                f"Заказчик: _______________ / {_short_fio(cp.get('director'))}")
    doc.add_paragraph().add_run("М.П.\t\t\t\t\t\t\t\tМ.П.")

    note = doc.add_paragraph()
    nr = note.add_run("\nПроект договора подготовлен автоматически. Перед "
                      "применением подлежит проверке и утверждению юристом.")
    nr.italic = True
    nr.font.size = Pt(8)
    nr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def _yr(sign_date):
    if isinstance(sign_date, str):
        try:
            return datetime.date.fromisoformat(sign_date).year
        except ValueError:
            return datetime.date.today().year
    if sign_date:
        return sign_date.year
    return datetime.date.today().year


def _short_fio(full):
    if not full:
        return "____________"
    parts = full.split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    return full
